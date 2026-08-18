"""Render a filled ITC short-form international sale contract to DOCX/PDF.

The operative wording follows the user-supplied 2010 ITC short-version model.
Only placeholders are filled and only user-selected model options are included.
No generative-AI output is inserted into legal clauses.
"""

from __future__ import annotations

from dataclasses import dataclass
from pathlib import Path

from .model import ContractData


@dataclass(frozen=True, slots=True)
class Block:
    text: str
    style: str = "body"  # title, label, h1, body, sub, note, signature


def _v(value: str, fallback: str = "[NOT SPECIFIED]") -> str:
    return value.strip() if value and value.strip() else fallback


def _payment_label(key: str) -> str:
    return {
        "advance": "Payment in advance",
        "collection": "Payment by documentary collection",
        "credit": "Payment by irrevocable documentary credit",
        "guarantee": "Payment backed by bank guarantee",
        "other": "Other payment arrangements",
    }[key]


def _doc_lines(c: ContractData) -> list[Block]:
    labels = {
        "commercial_invoice": "Commercial invoice",
        "transport_documents": "The following transport documents",
        "packing_list": "Packing list",
        "insurance_documents": "Insurance documents",
        "certificate_of_origin": "Certificate of origin",
        "certificate_of_inspection": "Certificate of inspection",
        "customs_documents": "Customs documents",
        "other_documents": "Other documents",
    }
    lines: list[Block] = []
    for key in c.documents:
        detail = ""
        if key == "transport_documents":
            detail = f": {_v(c.transport_document_details)}"
        elif key == "other_documents":
            detail = f": {_v(c.other_document_details)}"
        if c.document_copies:
            detail += f" (copies/details: {c.document_copies})"
        lines.append(Block(f"☒ {labels[key]}{detail}", "sub"))
    return lines


def build_contract_blocks(c: ContractData) -> list[Block]:
    """Return the contract as structured paragraphs after full validation."""
    c.validate()
    b: list[Block] = [
        Block("ITC MODEL CONTRACT FOR THE INTERNATIONAL COMMERCIAL SALE OF GOODS (SHORT VERSION)", "title"),
        Block("PARTIES:", "label"),
        Block("Seller", "label"),
        Block(f"Name (name of company): {_v(c.seller.name)}"),
        Block(f"Legal form (e.g. limited liability company): {_v(c.seller.legal_form)}"),
        Block(f"Country of incorporation and (if appropriate) trade register number: {_v(c.seller.country_and_register)}"),
        Block(f"Address (address of place of business of the seller, phone, fax, e-mail): {_v(c.seller.address)}"),
        Block(f"Represented by (surname and first name, address, position, legal title of representation): {_v(c.seller.represented_by)}"),
        Block("Buyer", "label"),
        Block(f"Name (name of company): {_v(c.buyer.name)}"),
        Block(f"Legal form (e.g. limited liability company): {_v(c.buyer.legal_form)}"),
        Block(f"Country of incorporation and (if appropriate) trade register number: {_v(c.buyer.country_and_register)}"),
        Block(f"Address (address of place of business of the buyer, phone, fax, e-mail): {_v(c.buyer.address)}"),
        Block(f"Represented by (surname and first name, address, position, legal title of representation): {_v(c.buyer.represented_by)}"),
        Block("Hereinafter: “the Parties”"),
        Block("OPERATIVE PROVISIONS", "label"),
        Block("1. Goods", "h1"),
        Block("1.1 Subject to the terms agreed in this contract, the Seller shall deliver the following good(s) (hereinafter: “the Goods”) to the Buyer."),
        Block(f"1.2 Description of the Goods (details necessary to define/specify the Goods which are the object of the sale, including required quality, description, certificates, country of origin, other details): {_v(c.goods_description)}"),
        Block("1.3 Quantity of the Goods (including unit of measurement)."),
        Block(f"1.3.1 Total quantity: {_v(c.total_quantity)}", "sub"),
        Block(f"1.3.2 Per delivery instalment: {_v(c.instalment_quantity, 'Not applicable')}", "sub"),
        Block(f"1.3.3 Tolerance percentage: Plus or minus {_v(c.tolerance_percentage, '0')} %", "sub"),
        Block(f"1.4 Inspection of the Goods (where an inspection is required, specify, as appropriate, details of organization responsible for inspecting quality and/or quantity, place and date and/or period of inspection, responsibility for inspection costs): {_v(c.inspection, 'No separate inspection agreed')}"),
        Block(f"1.5 Packaging: {_v(c.packaging)}"),
        Block(f"1.6 Other specification: {_v(c.other_specification, 'None')}"),
        Block("2. Delivery", "h1"),
        Block(f"2.1 Applicable International Chamber of Commerce (hereinafter: ICC) Incoterms (by reference to most recent version of the Incoterms at date of conclusion of the contract): {c.incoterm_rule} {_v(c.delivery_place)}, Incoterms® 2020."),
        Block(f"2.2 Place of delivery: {_v(c.delivery_place)}"),
        Block(f"2.3 Date or period of delivery: {_v(c.delivery_date_or_period)}"),
        Block(f"2.4 Carrier (name and address, where applicable): {_v(c.carrier, 'Not yet nominated / not applicable')}"),
        Block(f"2.5 Other delivery terms (if any): {_v(c.other_delivery_terms, 'None')}"),
        Block("3. Price", "h1"),
        Block(f"3.1 Total price: {_v(c.total_price)}"),
        Block(f"3.2 Price per unit of measurement (if appropriate): {_v(c.price_per_unit, 'Not applicable')}"),
        Block(f"3.3 Amount in numbers: {_v(c.amount_numbers, c.total_price)}"),
        Block(f"3.4 Amount in letters: {_v(c.amount_letters)}"),
        Block(f"3.5 Currency: {_v(c.currency)}"),
        Block(f"3.6 Method for determining the price (if appropriate): {_v(c.price_method, 'Fixed price')}"),
        Block("4. Payment conditions", "h1"),
        Block(f"4.1 Means of payment (e.g. cash, cheque, bank draft, transfer): {_v(c.means_of_payment)}"),
        Block(f"4.2 Details of Seller’s bank account (if appropriate): {_v(c.seller_bank_account, 'To be provided securely by the Seller')}"),
        Block(f"4.3 Time for payment: {_v(c.payment_time)}"),
        Block("The Parties may choose a payment arrangement among the possibilities set out below, in which case they should specify the arrangement chosen and provide the corresponding details:"),
        Block(f"☒ {_payment_label(c.payment_arrangement)}: {_v(c.payment_details)}", "sub"),
        Block("5. Documents", "h1"),
        Block("5.1 The Seller shall make available to the Buyer (or shall present to the bank specified by the Buyer) the following documents (tick corresponding boxes and indicate, as appropriate, the number of copies to be provided):"),
    ]
    b.extend(_doc_lines(c))
    b.extend([
        Block("5.2 In addition, the Seller shall make available to the Buyer the documents indicated in the ICC Incoterms the Parties have selected under Article 2 of this contract."),
        Block("6. Non-performance of the Buyer’s obligation to pay the price at the agreed time", "h1"),
        Block(f"6.1 If the Buyer fails to pay the price at the agreed time, the Seller shall fix to the Buyer an additional period of time of {_v(c.buyer_payment_additional_period)} for performance of payment. If the Buyer fails to pay the price at the expiration of the additional period, the Seller may declare this contract avoided in accordance with Article 10 of this contract."),
        Block(f"6.2 If the Buyer fails to pay the price at the agreed time, the Seller shall in any event be entitled, without limiting any other rights it may have, to charge interest on the outstanding amount (both before and after any judgment) at the rate of {_v(c.late_interest_rate, '0')} % per annum."),
        Block("7. Non-performance of the Seller’s obligation to deliver the Goods at the agreed time", "h1"),
        Block(f"7.1 If the Seller fails to deliver the Goods at the agreed time, the Buyer shall fix to the Seller an additional period of time of {_v(c.seller_delivery_additional_period)} for performance of delivery. If the Seller fails to deliver the Goods at the expiration of the additional period, the Buyer may declare this contract avoided in accordance with Article 10 of this contract."),
    ])
    if c.include_delay_liquidated_damages:
        b.append(Block(
            f"7.2 If the Seller is in delay in delivery of any goods as provided in this contract, the Buyer is entitled to claim liquidated damages equal to {_v(c.delay_daily_rate)}% of the price of those goods for each complete day of delay as from the agreed date of delivery or the last day of the agreed delivery period, as specified in Article 2 of this contract, provided the Buyer notifies the Seller of the delay. Where the Buyer so notifies the Seller within {_v(c.delay_notice_days)} days from the agreed date of delivery or the last day of the agreed delivery period, damages will run from the agreed date of delivery or from the last day of the agreed delivery period. Where the Buyer so notifies the Seller more than {_v(c.delay_notice_days)} days after the agreed date of delivery or the last day of the agreed delivery period, damages will run from the date of notice. Liquidated damages for delay shall not exceed {_v(c.delay_cap)}% of the price of the delayed goods. Liquidated damages for delay do not preclude avoidance of this contract in accordance with Article 10."
        ))
    b.extend([
        Block("8. Lack of conformity", "h1"),
        Block(f"8.1 The Buyer shall examine the Goods, or cause them to be examined within as short period as is practicable in the circumstances. The Buyer shall notify the Seller of any lack of conformity of the Goods, specifying the nature of the lack of conformity, within {_v(c.nonconformity_notice_days)} days after the Buyer has discovered or ought to have discovered the lack of conformity. In any event, the Buyer loses the right to rely on a lack of conformity if he fails to notify the Seller thereof at the latest within a period of {_v(c.nonconformity_long_stop)} from the date on which the Goods were actually handed over to the Buyer."),
        Block("8.2 Where the Buyer has given due notice of non-conformity to the Seller, the Buyer may at his option:"),
        Block("8.2.1 Require the Seller to deliver any missing quantity of the Goods, without any additional expense to the Buyer;", "sub"),
        Block("8.2.2 Require the Seller to replace the Goods with conforming goods, without any additional expense to the Buyer;", "sub"),
        Block("8.2.3 Require the Seller to repair the Goods, without any additional expense to the Buyer;", "sub"),
        Block("8.2.4 Reduce the price in the same proportion as the value that the Goods actually delivered had at the time of the delivery bears to the value that conforming goods would have had at that time. The Buyer may not reduce the price if the Seller replaces the Goods with conforming goods or repairs the Goods in accordance with paragraph 8.2.2 and 8.2.3 of this Article or if the Buyer refuses to accept such performance by the Seller;", "sub"),
        Block("8.2.5 Declare this contract avoided in accordance with Article 10 of this contract.", "sub"),
        Block("The Buyer shall in any event be entitled to claim damages."),
    ])
    if c.liability_limit:
        b.append(Block(f"8.3 The Seller’s liability under this Article for lack of conformity of the Goods is limited to {_v(c.liability_limit)}."))
    b.extend([
        Block("9. Transfer of property", "h1"),
    ])
    if c.retain_title:
        b.append(Block("The Seller must deliver to the Buyer the Goods specified in Article 1 of this contract free from any right or claim of a third person. The property in the Goods shall not pass to the Buyer until the Seller has received payment in full of the price of the Goods. Until property in the Goods passes to the Buyer, the Buyer shall keep the Goods separate from those of the Buyer and third parties and properly stored, protected and insured and identified as the Seller’s property."))
    else:
        b.append(Block("The Seller must deliver to the Buyer the Goods specified in Article 1 of this contract free from any right or claim of a third person."))
    b.extend([
        Block("10. Avoidance* of contract", "h1"),
        Block("10.1 There is a breach of contract where a party fails to perform any of its obligations under this contract, including defective, partial or late performance."),
        Block("10.2 There is a fundamental breach of contract where:"),
        Block("10.2.1 Strict compliance with the obligation which has not been performed is of the essence under this contract; or", "sub"),
        Block("10.2.2 The non-performance substantially deprives the aggrieved party of what it was reasonably entitled to expect under this contract.", "sub"),
    ])
    if c.additional_fundamental_breaches:
        b.append(Block(f"The Parties additionally agree that the following is to be considered as a fundamental breach of contract: {_v(c.additional_fundamental_breaches)}"))
    b.extend([
        Block(f"10.3 In a case of a breach of contract according to paragraph 10.1 of this Article, the aggrieved party shall, by notice to the other party, fix an additional period of time of {_v(c.general_breach_additional_period)} for performance. During the additional period of time the aggrieved party may withhold performance of its own reciprocal obligations and may claim damages, but may not declare this contract avoided. If the other party fails to perform its obligation within the additional period of time, the aggrieved party may declare this contract avoided."),
        Block("10.4 In case of a fundamental breach of contract according to paragraph 10.2 of this Article, the aggrieved party may declare this contract avoided without fixing an additional period of time for performance to the other party."),
        Block("10.5 A declaration of avoidance of this contract is effective only if made by notice to the other party."),
        Block("* Note: For the purposes of this Model Contract, the term “Avoidance” is taken from the CISG and means termination of contract.", "note"),
        Block("11. Force majeure – excuse for non-performance", "h1"),
        Block("11.1 “Force majeure” means war, emergency, accident, fire, earthquake, flood, storm, industrial strike or other impediment which the affected party proves was beyond its control and that it could not reasonably be expected to have taken the impediment into account at the time of the conclusion of this contract or to have avoided or overcome it or its consequences."),
        Block("11.2 A party affected by force majeure shall not be deemed to be in breach of this contract, or otherwise be liable to the other, by reason of any delay in performance, or the non-performance, of any of its obligations under this contract to the extent that the delay or non-performance is due to any force majeure of which it has notified the other party in accordance with Article 11.3. The time for performance of that obligation shall be extended accordingly, subject to Article 11.4."),
        Block("11.3 If any force majeure occurs in relation to either party which affects or is likely to affect the performance of any of its obligations under this contract, it shall notify the other party within a reasonable time as to the nature and extent of the circumstances in question and their effect on its ability to perform."),
    ])
    if c.force_majeure_negotiation_first:
        b.append(Block(f"11.4 If the performance by either party of any of its obligations under this contract is prevented or delayed by force majeure for a continuous period in excess of {_v(c.force_majeure_months)} months, the Parties shall negotiate in good faith, and use their best endeavours to agree upon such amendments to this contract or alternative arrangements as may be fair and reasonable with a view to alleviating its effects, but if they do not agree upon such amendments or arrangements within a further period of {_v(c.force_majeure_negotiation_days)} days, the other party shall be entitled to terminate this contract by giving written notice to the Party affected by the force majeure."))
    else:
        b.append(Block(f"11.4 If the performance by either party of any of its obligations under this contract is prevented or delayed by force majeure for a continuous period in excess of {_v(c.force_majeure_months)} months, the other party shall be entitled to terminate this contract by giving written notice to the Party affected by the force majeure."))
    entire = "12.1 This contract sets out the entire agreement between the Parties. Neither party has entered into this contract in reliance upon any representation, warranty or undertaking of the other party that is not expressly set out or referred to in this contract. This Article shall not exclude any liability for fraudulent misrepresentation."
    if c.supersedes_previous:
        entire += " This contract supersedes any previous agreement or understanding relating its subject matter."
    b.extend([
        Block("12. Entire agreement", "h1"),
        Block(entire),
        Block("12.2 This contract may not be varied except by an agreement of the Parties in writing (which may include e-mail)."),
        Block("13. Notices", "h1"),
        Block("13.1 Any notice under this contract shall be in writing (which may include e-mail) and may be served by leaving it or sending it to the address of the other party as specified in Article 13.2 below, in a manner that ensures receipt of the notice can be proved."),
        Block("13.2 For the purposes of Article 13.1, notification details are the following, unless other details have been duly notified in accordance with this Article:"),
        Block(f"– Seller: {_v(c.seller_notice_details, c.seller.address)};", "sub"),
        Block(f"– Buyer: {_v(c.buyer_notice_details, c.buyer.address)}.", "sub"),
        Block("14. Dispute resolution procedure", "h1"),
    ])
    if c.dispute_method == "institutional":
        b.append(Block(f"Any dispute, controversy or claim arising out of or relating to this contract, including its conclusion, interpretation, performance, breach, termination or invalidity, shall be finally settled under the rules of {_v(c.arbitration_institution)} by {_v(c.arbitrators)} appointed in accordance with the said rules. The place of arbitration shall be {_v(c.arbitration_place)}. The language of the arbitration shall be {_v(c.arbitration_language)}."))
    elif c.dispute_method == "ad_hoc":
        b.append(Block(f"Any dispute, controversy or claim arising out of or relating to this contract, including its conclusion, interpretation, performance, breach, termination or invalidity, shall be finally settled under the rules of UNCITRAL by {_v(c.arbitrators)} appointed by {_v(c.appointing_authority)}. The place of arbitration shall be {_v(c.arbitration_place)}. The language of the arbitration shall be {_v(c.arbitration_language)}."))
    else:
        b.append(Block(f"Any dispute, controversy or claim arising out of or relating to this contract, in particular its conclusion, interpretation, performance, breach, termination or invalidity, shall be finally settled by the courts of {_v(c.court_place_country)} which will have exclusive jurisdiction."))
    b.extend([
        Block("15. Applicable law and guiding principles", "h1"),
        Block("15.1 Questions relating to this contract that are not settled by the provisions contained in the contract itself shall be governed by the United Nations Convention on Contracts for the International Sale of Goods (Vienna Sales Convention of 1980, hereafter referred to as CISG)."),
        Block(f"Questions not covered by the CISG shall be governed by the UNIDROIT Principles of International Commercial Contracts (hereafter referred to as UNIDROIT Principles), and to the extent that such questions are not covered by the UNIDROIT Principles, by reference to the applicable national law of {_v(c.national_law)}."),
        Block("15.2 This contract shall be performed in a spirit of good faith and fair dealing."),
        Block("DATE AND SIGNATURE OF THE PARTIES", "label"),
        Block(f"Date: {_v(c.contract_date)}", "signature"),
        Block(f"Seller – Name: {_v(c.seller_signatory_name, c.seller.represented_by)}", "signature"),
        Block("Signature: ____________________________________", "signature"),
        Block(f"Buyer – Name: {_v(c.buyer_signatory_name, c.buyer.represented_by)}", "signature"),
        Block("Signature: ____________________________________", "signature"),
    ])
    return b


def render_docx(c: ContractData, output_path: str | Path) -> Path:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Mm, Pt

    path = Path(output_path)
    path.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    sec = doc.sections[0]
    sec.page_width, sec.page_height = Mm(210), Mm(297)
    sec.top_margin = sec.bottom_margin = Mm(18)
    sec.left_margin = sec.right_margin = Mm(23)
    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
    normal.font.size = Pt(10)
    normal.paragraph_format.space_after = Pt(4)
    normal.paragraph_format.line_spacing = 1.0

    for block in build_contract_blocks(c):
        if block.text == "DATE AND SIGNATURE OF THE PARTIES":
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(9)
            p.paragraph_format.space_after = Pt(5)
            run = p.add_run(block.text)
            run.bold = True

            table = doc.add_table(rows=4, cols=2)
            table.autofit = False
            for row in table.rows:
                row.cells[0].width = Mm(80)
                row.cells[1].width = Mm(80)
            values = [
                ("Seller", "Buyer"),
                (f"Date: {_v(c.contract_date)}", f"Date: {_v(c.contract_date)}"),
                (f"Name: {_v(c.seller_signatory_name, c.seller.represented_by)}", f"Name: {_v(c.buyer_signatory_name, c.buyer.represented_by)}"),
                ("Signature: __________________________", "Signature: __________________________"),
            ]
            for row_index, (row, pair) in enumerate(zip(table.rows, values)):
                for cell, text in zip(row.cells, pair):
                    cell.text = text
                    for para in cell.paragraphs:
                        para.paragraph_format.space_after = Pt(4)
                    if row_index == 0:
                        cell.paragraphs[0].runs[0].bold = True
            tbl_pr = table._tbl.tblPr
            borders = OxmlElement("w:tblBorders")
            for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
                tag = OxmlElement(f"w:{edge}")
                tag.set(qn("w:val"), "nil")
                borders.append(tag)
            tbl_pr.append(borders)
            break
        p = doc.add_paragraph()
        p.paragraph_format.keep_together = True
        if block.style == "title":
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_after = Pt(16)
            run = p.add_run(block.text)
            run.bold, run.font.size = True, Pt(13)
        elif block.style == "h1":
            p.paragraph_format.keep_with_next = True
            p.paragraph_format.space_before = Pt(9)
            p.paragraph_format.space_after = Pt(4)
            run = p.add_run(block.text)
            run.bold, run.font.size = True, Pt(11)
        elif block.style == "label":
            p.paragraph_format.space_before = Pt(7)
            run = p.add_run(block.text)
            run.bold = True
        elif block.style == "sub":
            p.paragraph_format.left_indent = Mm(8)
            p.add_run(block.text)
        elif block.style == "note":
            run = p.add_run(block.text)
            run.italic, run.font.size = True, Pt(8.5)
        elif block.style == "signature":
            p.paragraph_format.space_after = Pt(10)
            p.add_run(block.text)
        else:
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p.add_run(block.text)
    props = doc.core_properties
    props.title = "International Commercial Sale of Goods (ITC short-form model)"
    props.subject = "Ready4Trade draft generated from user-provided data"
    props.author = "Ready4Trade"
    doc.save(path)
    return path


def render_pdf(c: ContractData, output_path: str | Path) -> Path:
    from reportlab.lib.enums import TA_CENTER, TA_LEFT
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
    from reportlab.lib.units import mm
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont
    from reportlab.platypus import BaseDocTemplate, Frame, PageTemplate, Paragraph, Spacer
    from xml.sax.saxutils import escape

    path = Path(output_path)
    path.parent.mkdir(parents=True, exist_ok=True)
    regular, bold, italic = "Times-Roman", "Times-Bold", "Times-Italic"
    font_sets = (
        (
            "/usr/share/fonts/truetype/dejavu/DejaVuSerif.ttf",
            "/usr/share/fonts/truetype/dejavu/DejaVuSerif-Bold.ttf",
            "/usr/share/fonts/truetype/dejavu/DejaVuSerif-Italic.ttf",
        ),
        (
            "/usr/share/fonts/truetype/liberation2/LiberationSerif-Regular.ttf",
            "/usr/share/fonts/truetype/liberation2/LiberationSerif-Bold.ttf",
            "/usr/share/fonts/truetype/liberation2/LiberationSerif-Italic.ttf",
        ),
    )
    for normal_path, bold_path, italic_path in font_sets:
        if all(Path(item).exists() for item in (normal_path, bold_path, italic_path)):
            pdfmetrics.registerFont(TTFont("R4TSerif", normal_path))
            pdfmetrics.registerFont(TTFont("R4TSerif-Bold", bold_path))
            pdfmetrics.registerFont(TTFont("R4TSerif-Italic", italic_path))
            regular, bold, italic = "R4TSerif", "R4TSerif-Bold", "R4TSerif-Italic"
            break

    styles = getSampleStyleSheet()
    body = ParagraphStyle("ContractBody", parent=styles["BodyText"], fontName=regular, fontSize=9.5, leading=11.2, alignment=TA_LEFT, spaceAfter=4)
    title = ParagraphStyle("ContractTitle", parent=body, fontName=bold, fontSize=12.5, leading=14, alignment=TA_CENTER, spaceAfter=14)
    h1 = ParagraphStyle("ContractH1", parent=body, fontName=bold, fontSize=10.5, leading=12, spaceBefore=7, spaceAfter=3, keepWithNext=True)
    label = ParagraphStyle("ContractLabel", parent=body, fontName=bold, spaceBefore=6)
    sub = ParagraphStyle("ContractSub", parent=body, leftIndent=8 * mm)
    note = ParagraphStyle("ContractNote", parent=body, fontName=italic, fontSize=8, leading=9.5)
    sig = ParagraphStyle("ContractSig", parent=body, spaceAfter=9)
    style_map = {"title": title, "h1": h1, "label": label, "sub": sub, "note": note, "signature": sig, "body": body}

    doc = BaseDocTemplate(str(path), pagesize=A4, leftMargin=23 * mm, rightMargin=23 * mm, topMargin=18 * mm, bottomMargin=18 * mm, title="International Commercial Sale of Goods", author="Ready4Trade")
    frame = Frame(doc.leftMargin, doc.bottomMargin, doc.width, doc.height, id="contract")

    def footer(canvas, document):
        canvas.saveState()
        canvas.setFont(regular, 8)
        canvas.drawRightString(A4[0] - 23 * mm, 9 * mm, str(document.page))
        canvas.restoreState()

    doc.addPageTemplates([PageTemplate(id="contract-pages", frames=[frame], onPage=footer)])
    story = []
    for block in build_contract_blocks(c):
        text = escape(block.text.replace("☒", "[X]")).replace("\n", "<br/>")
        story.append(Paragraph(text, style_map[block.style]))
    doc.build(story)
    return path


def render_both(c: ContractData, directory: str | Path, stem: str = "international_sale_contract") -> tuple[Path, Path]:
    directory = Path(directory)
    directory.mkdir(parents=True, exist_ok=True)
    return render_docx(c, directory / f"{stem}.docx"), render_pdf(c, directory / f"{stem}.pdf")
